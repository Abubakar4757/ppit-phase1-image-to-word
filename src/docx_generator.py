from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from src.formatting_detector import FormattedBlock


class DocxGenerator:
    def __init__(self) -> None:
        pass

    def generate_docx(self, blocks: List[FormattedBlock], output_path: str | Path) -> bool:
        doc = Document()
        
        # 1. Page Setup
        section = doc.sections[0]
        section.page_height = Inches(11.69) # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # 2. Add Blocks
        for block in blocks:
            if not block.text.strip():
                continue
                
            p = doc.add_paragraph()
            
            # Map block types to styles
            if block.block_type == "heading":
                p.style = doc.styles["Heading 1"]
            elif block.block_type == "bullet":
                p.style = doc.styles["List Bullet"]
            elif block.block_type == "numbered":
                p.style = doc.styles["List Number"]
            else:
                p.style = doc.styles["Normal"]
                
            # Apply alignment
            if block.alignment == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif block.alignment == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            # Apply indentation
            if block.indent_level > 0:
                p.paragraph_format.left_indent = Pt(18 * block.indent_level)
                
            # Add text with runs for bold/italic
            run = p.add_run(block.text)
            if block.is_bold:
                run.bold = True
            if block.is_italic:
                run.italic = True
                
        # 3. Add Metadata (optional footer or at end)
        doc.add_page_break()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer_run.font.size = Pt(8)
        footer_run.italic = True
        
        # 4. Save
        try:
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            return True
        except Exception as e:
            print(f"Error saving docx: {e}")
            return False


def generate_docx(blocks: List[FormattedBlock], output_path: str | Path) -> bool:
    return DocxGenerator().generate_docx(blocks, output_path)
